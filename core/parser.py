from pathlib import Path


def extract_text(file_path: str) -> str:
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")
    if not path.is_file():
        raise ValueError(f"路径不是文件: {file_path}")

    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(path)
    elif suffix == ".docx":
        return _extract_docx(path)
    else:
        raise ValueError(f"不支持的文件格式: {suffix}，仅支持 PDF 和 DOCX")


def _extract_pdf(path: Path) -> str:
    import fitz

    try:
        doc = fitz.open(str(path))
    except Exception as e:
        raise ValueError(f"文件无法读取，可能已损坏: {e}")

    text_parts = []
    for page in doc:
        text = page.get_text()
        if text:
            text_parts.append(text.strip())
    doc.close()

    result = "\n".join(text_parts)
    if len(result) < 50:
        raise ValueError(
            "提取的文本过短（<50字符）。请确认上传的是文本型PDF，图片型PDF无法提取文字。"
        )
    return result


def _extract_docx(path: Path) -> str:
    from docx import Document

    try:
        doc = Document(str(path))
    except Exception as e:
        raise ValueError(f"文件无法读取，可能已损坏: {e}")

    text_parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            text_parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    result = "\n".join(text_parts)
    if not result:
        raise ValueError("文档内容为空")
    return result
